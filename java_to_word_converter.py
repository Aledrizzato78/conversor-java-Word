import os
from docx import Document
from docx.shared import Pt

def processar_arquivo(file_path, documento):
    _, file_name = os.path.split(file_path)
    documento.add_heading(f'Arquivo: {file_name}', level=2)
    
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            conteudo = file.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                conteudo = file.read()
        except Exception as e:
            documento.add_paragraph(f"Erro ao ler o arquivo: {str(e)}")
            return
    except Exception as e:
        documento.add_paragraph(f"Erro ao processar o arquivo: {str(e)}")
        return
    
    p = documento.add_paragraph()
    run = p.add_run(conteudo)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

def percorrer_diretorios(root_dir, documento):
    for subdir, dirs, files in os.walk(root_dir):
        rel_path = os.path.relpath(subdir, root_dir)
        documento.add_heading(f'Diretório: {rel_path}', level=1)
        
        for file in files:
            if file.endswith(('.java', '.xml', '.properties', '.gradle', '.md')):
                file_path = os.path.join(subdir, file)
                processar_arquivo(file_path, documento)

def main():
    root_dir = 'C:/Users/Latitude 5280/IdeaProjects/podoc'
    documento = Document()
    documento.add_heading('Projeto Java: podoc', level=0)

    percorrer_diretorios(root_dir, documento)

    output_path = 'G:/Meu Drive/Projetos/podoc.docx'
    try:
        documento.save(output_path)
        print(f"Arquivo Word criado com sucesso em {output_path}")
    except Exception as e:
        print(f"Erro ao salvar o arquivo Word: {str(e)}")

if __name__ == "__main__":
    main()